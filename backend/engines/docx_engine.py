"""
VERITAS Docs — DOCX Engine
============================
Renders a VeritasDocument to a branded Word document using python-docx.
Features: branded heading styles, header/footer with Ω, alternating-row tables.
"""

import os
import re
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

import sys
sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
from models import BrandPreset, VeritasDocument


def _hex_to_rgb(hex_color: str) -> RGBColor:
    """Convert hex color string to docx RGBColor."""
    h = hex_color.lstrip('#')
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _add_header_footer(doc: Document, brand: BrandPreset, title: str):
    """Add branded header and footer to the document."""
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # Header
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run(f"{brand.logo_text} VERITAS DOCS")
    run.font.size = Pt(8)
    run.font.color.rgb = _hex_to_rgb(brand.gold_dim)
    run.font.bold = True

    # Footer
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run(brand.footer_text)
    run.font.size = Pt(8)
    run.font.color.rgb = _hex_to_rgb('#999999')


def _parse_inline(text: str) -> list:
    """Parse inline markdown and return list of (text, bold, italic) tuples."""
    segments = []
    # Simple parser: split on **bold** and *italic* markers
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            segments.append((part[2:-2], True, False))
        elif part.startswith('*') and part.endswith('*'):
            segments.append((part[1:-1], False, True))
        else:
            # Strip markdown links
            clean = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', part)
            if clean:
                segments.append((clean, False, False))
    return segments


def render(doc_data: VeritasDocument, brand: BrandPreset, output_path: Path) -> Path:
    """Render a VeritasDocument to a branded DOCX file."""
    document = Document()
    gold = _hex_to_rgb(brand.gold)
    obsidian = _hex_to_rgb(brand.obsidian)

    _add_header_footer(document, brand, doc_data.title)

    # Cover Page
    document.add_paragraph()  # spacer
    cover = document.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = cover.add_run(brand.logo_text)
    run.font.size = Pt(48)
    run.font.color.rgb = gold
    run.font.bold = True

    brand_line = document.add_paragraph()
    brand_line.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = brand_line.add_run("VERITAS DOCS")
    run.font.size = Pt(14)
    run.font.color.rgb = gold
    run.font.bold = True

    # Title
    title_p = document.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title_p.add_run(doc_data.title)
    run.font.size = Pt(24)
    run.font.color.rgb = gold
    run.font.bold = True

    if doc_data.subtitle:
        sub_p = document.add_paragraph()
        run = sub_p.add_run(doc_data.subtitle)
        run.font.size = Pt(12)
        run.font.color.rgb = _hex_to_rgb('#AAAAAA')

    # Date
    date_p = document.add_paragraph()
    run = date_p.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
    run.font.size = Pt(10)
    run.font.color.rgb = _hex_to_rgb('#999999')

    # Confidentiality
    conf = brand.confidentiality_label
    if conf:
        conf_p = document.add_paragraph()
        run = conf_p.add_run(conf)
        run.font.size = Pt(9)
        run.font.color.rgb = _hex_to_rgb(brand.confidentiality_color)
        run.font.bold = True

    # Motto
    document.add_paragraph()
    motto_p = document.add_paragraph()
    run = motto_p.add_run(brand.motto)
    run.font.size = Pt(10)
    run.font.color.rgb = gold
    run.font.italic = True

    document.add_page_break()

    # Parse markdown body
    lines = doc_data.content.split('\n')
    in_code_block = False
    code_buffer = []
    table_buffer = []
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Code blocks
        if stripped.startswith('```'):
            if in_code_block:
                code_text = '\n'.join(code_buffer)
                p = document.add_paragraph()
                run = p.add_run(code_text)
                run.font.name = 'Courier New'
                run.font.size = Pt(8)
                run.font.color.rgb = _hex_to_rgb('#444444')
                code_buffer = []
                in_code_block = False
            else:
                # Flush table
                if table_buffer:
                    _add_table(document, table_buffer, brand)
                    table_buffer = []
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue

        if not stripped:
            if table_buffer:
                _add_table(document, table_buffer, brand)
                table_buffer = []
            i += 1
            continue

        # Table rows
        if '|' in stripped and stripped.startswith('|'):
            cells = [c.strip() for c in stripped.split('|')[1:-1]]
            if all(re.match(r'^[-:]+$', c) for c in cells):
                i += 1
                continue
            table_buffer.append(cells)
            i += 1
            continue
        elif table_buffer:
            _add_table(document, table_buffer, brand)
            table_buffer = []

        # Headings
        if stripped.startswith('# ') and not stripped.startswith('## '):
            h = document.add_heading(stripped[2:].strip(), level=1)
            for run in h.runs:
                run.font.color.rgb = gold
            i += 1
            continue

        if stripped.startswith('## '):
            h = document.add_heading(stripped[3:].strip(), level=2)
            for run in h.runs:
                run.font.color.rgb = gold
            i += 1
            continue

        if stripped.startswith('### '):
            h = document.add_heading(stripped[4:].strip(), level=3)
            i += 1
            continue

        # Horizontal rule / page break
        if re.match(r'^---+$', stripped):
            document.add_page_break()
            i += 1
            continue

        # Blockquote
        if stripped.startswith('>'):
            p = document.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            run = p.add_run(stripped[1:].strip())
            run.font.italic = True
            run.font.color.rgb = gold
            i += 1
            continue

        # Bullets
        if stripped.startswith('- ') or stripped.startswith('* '):
            p = document.add_paragraph(style='List Bullet')
            segments = _parse_inline(stripped[2:].strip())
            for text, bold, italic in segments:
                run = p.add_run(text)
                run.font.bold = bold
                run.font.italic = italic
                run.font.size = Pt(10)
            i += 1
            continue

        # Body text
        p = document.add_paragraph()
        segments = _parse_inline(stripped)
        for text, bold, italic in segments:
            run = p.add_run(text)
            run.font.bold = bold
            run.font.italic = italic
            run.font.size = Pt(10)
        i += 1

    # Flush remaining
    if table_buffer:
        _add_table(document, table_buffer, brand)
    if code_buffer:
        p = document.add_paragraph()
        run = p.add_run('\n'.join(code_buffer))
        run.font.name = 'Courier New'
        run.font.size = Pt(8)

    # Footer motto
    document.add_paragraph()
    final = document.add_paragraph()
    final.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = final.add_run(f"{brand.logo_text} {brand.motto}")
    run.font.size = Pt(9)
    run.font.color.rgb = gold
    run.font.italic = True

    os.makedirs(output_path.parent, exist_ok=True)
    document.save(str(output_path))
    return output_path


def _add_table(document: Document, rows: list, brand: BrandPreset):
    """Add a branded table to the document."""
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    table = document.add_table(rows=len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for ri, row in enumerate(rows):
        for ci, cell_text in enumerate(row):
            if ci < n_cols:
                cell = table.rows[ri].cells[ci]
                cell.text = cell_text
                p = cell.paragraphs[0]
                for run in p.runs:
                    run.font.size = Pt(9)
                    if ri == 0:
                        run.font.bold = True

    document.add_paragraph()  # spacer after table
